"""Estrazione testo locale, gratuita, per i formati "digitali" (PDF con
strato di testo, DOCX, XLSX). Nessuna chiamata LLM qui - solo parsing.

Per i PDF: usata anche come pre-check a costo zero per decidere il
routing in ingest_documento.py - se un PDF ha gia' abbastanza testo
estraibile, non serve passare dalla visione Sonnet (vedi DECISIONS.md,
Tappa 5). Se il testo estratto e' trascurabile (scansione), il chiamante
passa al percorso visivo.
"""
from __future__ import annotations

import io

import openpyxl
import pypdf
from docx import Document

# Sotto questa soglia di caratteri un PDF e' considerato scansionato
# (nessuno strato di testo utile), non un fallimento di pypdf su un PDF
# comunque testuale ma corto: un documento reale (fattura, contratto) ha
# sempre almeno qualche riga di testo utile.
SOGLIA_TESTO_DIGITALE = 40


def estrai_testo_pdf(contenuto: bytes) -> str:
    """Estrae il testo nativo di un PDF (nessun OCR). Stringa vuota o
    corta se il PDF non ha uno strato di testo (scansione) o se il file
    non è un PDF valido — trattato come "nessun testo digitale", non un
    crash: il chiamante passa al percorso visivo, che gestisce l'errore
    in modo esplicito se il file è davvero illeggibile."""
    try:
        lettore = pypdf.PdfReader(io.BytesIO(contenuto))
        pagine = [pagina.extract_text() or "" for pagina in lettore.pages]
    except Exception:
        return ""
    return "\n\n".join(pagine).strip()


def pdf_ha_testo_digitale(contenuto: bytes) -> bool:
    return len(estrai_testo_pdf(contenuto)) >= SOGLIA_TESTO_DIGITALE


def numero_pagine_pdf(contenuto: bytes) -> int:
    return len(pypdf.PdfReader(io.BytesIO(contenuto)).pages)


def pdf_e_cifrato(contenuto: bytes) -> bool:
    """Un PDF protetto da password va rifiutato con un messaggio chiaro
    PRIMA del routing: pypdf estrarrebbe stringa vuota (=> "scansione") e
    l'API rifiuterebbe il documento cifrato con un errore grezzo. Un file
    corrotto NON e' "cifrato": segue il percorso normale."""
    try:
        return bool(pypdf.PdfReader(io.BytesIO(contenuto)).is_encrypted)
    except Exception:
        return False


# Sotto questa soglia di caratteri una SINGOLA pagina è considerata senza
# testo utile (cfr. SOGLIA_TESTO_DIGITALE che vale sul totale).
SOGLIA_TESTO_PAGINA = 25


def indici_pagine_scansione(contenuto: bytes) -> list[int]:
    """Pagine senza testo utile MA con immagini: quasi certamente pagine
    scansionate dentro un PDF altrimenti digitale (es. contratto con
    copertina digitale + allegati scannerizzati). Il routing per soglia sul
    testo TOTALE le classificherebbe "digitali" perdendole in silenzio -
    trappola trovata rivalutando la Tappa 5. Una pagina bianca (niente
    testo, niente immagini) non conta: è comune e innocua."""
    try:
        lettore = pypdf.PdfReader(io.BytesIO(contenuto))
    except Exception:
        return []
    indici = []
    for indice, pagina in enumerate(lettore.pages):
        try:
            testo = (pagina.extract_text() or "").strip()
        except Exception:
            testo = ""
        if len(testo) >= SOGLIA_TESTO_PAGINA:
            continue
        try:
            ha_immagini = bool(pagina.images)
        except Exception:
            ha_immagini = False
        if ha_immagini:
            indici.append(indice)
    return indici


def estrai_testo_docx(contenuto: bytes) -> str:
    documento = Document(io.BytesIO(contenuto))
    paragrafi = [p.text for p in documento.paragraphs if p.text.strip()]
    for tabella in documento.tables:
        for riga in tabella.rows:
            testo_riga = " | ".join(cella.text for cella in riga.cells)
            if testo_riga.strip():
                paragrafi.append(testo_riga)
    return "\n".join(paragrafi).strip()


def estrai_testo_xlsx(contenuto: bytes) -> str:
    cartella = openpyxl.load_workbook(io.BytesIO(contenuto), data_only=True)
    fogli_testo = []
    for foglio in cartella.worksheets:
        righe_testo = []
        for riga in foglio.iter_rows(values_only=True):
            valori = [str(v) for v in riga if v is not None]
            if valori:
                righe_testo.append(", ".join(valori))
        if righe_testo:
            fogli_testo.append(f"[{foglio.title}]\n" + "\n".join(righe_testo))
    return "\n\n".join(fogli_testo).strip()
